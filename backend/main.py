from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
import json
import uuid
from pathlib import Path

app = FastAPI()

# Directories for storing templates and generated documents.  Use paths relative
# to this file so the API functions even when the working directory changes.
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / 'templates'
OUTPUT_DIR = BASE_DIR / 'generated'


# ensure storage directories exist
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

class TemplateRequest(BaseModel):
    """Request body for creating a template."""

    style_desc: str = ''
    hint: str = 'minimal'


def _apply_hint_style(doc: Document, hint: str) -> None:
    """Apply basic styling to the document based on the hint.

    This is a very small demo of what could be a much more powerful styling
    engine.  The goal is simply to make the produced template visually reflect
    the selected hint so the user sees that something has happened when the
    Create button is pressed.
    """

    style = doc.styles['Normal'].font
    if hint == 'corporate':
        style.name = 'Calibri'
        style.size = Pt(12)
    elif hint == 'modern':
        style.name = 'Arial'
        style.size = Pt(11)
    elif hint == 'classic':
        style.name = 'Times New Roman'
        style.size = Pt(12)
    else:  # minimal
        style.name = 'Calibri'
        style.size = Pt(11)



class TemplateRequest(BaseModel):
    """Request body for creating a template."""

    style_desc: str = ''
    hint: str = 'minimal'


def _apply_hint_style(doc: Document, hint: str) -> None:
    """Apply basic styling to the document based on the hint.

    This is a very small demo of what could be a much more powerful styling
    engine.  The goal is simply to make the produced template visually reflect
    the selected hint so the user sees that something has happened when the
    Create button is pressed.
    """

    style = doc.styles['Normal'].font
    if hint == 'corporate':
        style.name = 'Calibri'
        style.size = Pt(12)
    elif hint == 'modern':
        style.name = 'Arial'
        style.size = Pt(11)
    elif hint == 'classic':
        style.name = 'Times New Roman'
        style.size = Pt(12)
    else:  # minimal
        style.name = 'Calibri'
        style.size = Pt(11)



@app.post('/templates')
async def create_template(req: TemplateRequest):
    """Create a Word template with a very simple style applied.

    The style description and hint roughly determine the fonts used in the
    generated template so the user can immediately see the effect.
    """

    template_path = TEMPLATE_DIR / f"{uuid.uuid4()}.docx"
    doc = Document()
    _apply_hint_style(doc, req.hint)

    # Use the style description as a heading if supplied
    doc.add_heading(req.style_desc or 'Document Title', level=1)

    # Include a simple placeholder so the template can be rendered later
    doc.add_paragraph('Hello {{ name }}!')


    try:
        doc.save(template_path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f'Failed to save template: {exc}')

    template_id = template_path.stem
    return {
        'template_id': template_id,
        'download_url': f'/templates/{template_id}',
    }


@app.get('/templates/{template_id}')
async def download_template(template_id: str):
    """Return the generated template for download."""
    template_path = TEMPLATE_DIR / f'{template_id}.docx'
    if not template_path.exists():
        raise HTTPException(status_code=404, detail='Template not found')
    return FileResponse(template_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f'{template_id}.docx')

@app.post('/documents/{template_id}')
async def generate_document(template_id: str, data: UploadFile = File(...)):
    template_path = TEMPLATE_DIR / f'{template_id}.docx'
    if not template_path.exists():
        raise HTTPException(status_code=404, detail='Template not found')
    context = json.loads(await data.read())
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    out_path = OUTPUT_DIR / f"{template_id}-{uuid.uuid4()}.docx"
    tpl.save(out_path)
    return {'document': str(out_path)}
